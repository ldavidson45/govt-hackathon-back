from docx import Document
from django.http import JsonResponse


def parse_document(document):

    doc = Document(document)

    max_search = 30

    document_find = ['Description', "CONTRACTOR PERSONNEL", "General", "Combined Synops",
                    "Line items", "GOVERN", "Applicable", "RESPONSES", "PURPOSE", "BACKGROUND", "PWS", "Applicable Provisions", "Offer Eval", "Procedure",
                    "Submission", "BACKGROUND", "REQUIRED SERVICES", "ROOM & BOARD", "LAUNDRY", "DIRECT SERVICES",
                    "Care Planning", "Recovery", "Medication Management", "Case Management", "Treatment", "Group Activities",
                    "Discharge Planning", "VA Coordination", "Transportation", "ADMINISTRATIVE", "COMMUNICATIONS",
                    "ADMISSIONS", "DOCUMENTATION", "ABSENCES", "CRITICAL", "Scope", "Objectives", "Tree Services",
                    "WORK CLEARANCE", "FACILITY", "GENERAL REQUIREMENTS", "INSPECTION", "QUALITY"
                    ]

    main_dictionary = {}
    Default = "No relevant information found"
    compiled_data_list = []
    for i in range(len(doc.paragraphs)):
        for x in range(len(document_find)):
            if document_find[x] in doc.paragraphs[i].text[:max_search]:
                count = 0
                test = 0
                while(not doc.paragraphs[count+i].text.strip()):
                    count += 1
                if(count > 7 or count < 1):
                    main_dictionary[document_find[x]] = doc.paragraphs[i+2].text
                else:
                    main_dictionary[document_find[x]] = doc.paragraphs[i+count].text

        return main_dictionary
